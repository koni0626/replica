"""
Markdown(.md/.markdown) → Word(.docx) 変換ツール

優先ルート:
- pypandoc + pandoc が利用可能ならそれを使って高品質に変換（参照テンプレートで游明朝に設定）
- 利用不可の場合は python-docx を用いた簡易レンダラで見出し/段落/箇条書き/番号付き/コードブロック/強調(太字/斜体)の最小対応

依存:
- 推奨: pypandoc (外部に pandoc コマンドが必要)
- フォールバック: python-docx
    pip install python-docx

公開API:
- convert_md_to_docx(md_path, docx_path, encoding="utf-8", ensure_parent=True) -> tuple[str, int]
    返り値: (出力先絶対パス, 行数/段落数の目安)
"""
from __future__ import annotations

from pathlib import Path
import os
import re
import tempfile
from typing import Tuple

# pypandoc は任意依存
try:
    import pypandoc  # type: ignore
except Exception:
    pypandoc = None  # type: ignore

try:
    from docx import Document  # python-docx
    from docx.shared import Pt
    from docx.oxml.ns import qn
except Exception:
    Document = None  # type: ignore
    qn = None  # type: ignore


class MarkdownToDocxError(Exception):
    pass


def _ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def _has_md_ext(p: Path) -> bool:
    return p.suffix.lower() in {".md", ".markdown"}


def _build_reference_docx(font_name: str = "游明朝") -> Path | None:
    """游明朝を既定スタイルに設定した参照テンプレート(.docx)を一時生成して返す。
    python-docx が無い場合は None。
    """
    if Document is None or qn is None:
        return None
    try:
        doc = Document()
        styles = doc.styles
        target_styles = [
            "Normal",
            "Heading 1", "Heading 2", "Heading 3",
            "Heading 4", "Heading 5", "Heading 6",
            "List Bullet", "List Number",
        ]
        for sname in target_styles:
            try:
                st = styles[sname]
            except Exception:
                continue
            # 欧文フォント
            st.font.name = font_name
            # 東アジアフォント（日本語）
            try:
                if st._element is not None:
                    rPr = getattr(st._element, 'rPr', None)
                    if rPr is None:
                        rPr = st._element.get_or_add_rPr()
                    rFonts = getattr(rPr, 'rFonts', None)
                    if rFonts is None:
                        rFonts = rPr.get_or_add_rFonts()
                    rFonts.set(qn('w:eastAsia'), font_name)
            except Exception:
                pass
        # 参考になる文字列を軽く入れておく（必須ではない）
        doc.add_paragraph("テンプレート: 游明朝（Normal）")
        tmp = Path(tempfile.gettempdir()) / "reference_yumincho.docx"
        doc.save(str(tmp))
        return tmp
    except Exception:
        return None


def convert_md_to_docx(md_path: str, docx_path: str, encoding: str = "utf-8", ensure_parent: bool = True) -> Tuple[str, int]:
    """Markdown を Word(.docx) に変換して保存。

    pypandoc が利用可能なら優先して使用し、不可なら簡易レンダリングにフォールバックする。

    Returns: (出力絶対パス, 段落数の目安)
    """
    src = Path(md_path).expanduser().resolve()
    dst = Path(docx_path).expanduser().resolve()

    if not _has_md_ext(src):
        raise MarkdownToDocxError(f"Unsupported extension for markdown: {src}")
    if src.is_dir() or (not src.exists()):
        raise MarkdownToDocxError(f"Markdown not found: {src}")
    if ensure_parent:
        _ensure_parent(dst)

    # 1) pypandoc 経由（可能なら）
    if pypandoc is not None:
        try:
            extra_args = []
            # 参照テンプレート（游明朝）を生成できる場合は指定
            ref_doc = _build_reference_docx(font_name="游明朝")
            if ref_doc is not None and ref_doc.exists():
                extra_args.append(f"--reference-doc={ref_doc}")
            # GitHub Flavored Markdown として処理（表の互換性が高い）
            pypandoc.convert_file(
                str(src),
                to='docx',
                format='gfm',
                outputfile=str(dst),
                extra_args=extra_args,
            )
            return (str(dst), 0)
        except Exception:
            # フォールバックへ
            pass

    # 2) 簡易レンダリング（python-docx 必須）
    if Document is None:
        raise MarkdownToDocxError("Neither pypandoc nor python-docx is available. Install pypandoc+pandoc or python-docx.")

    text = src.read_text(encoding=encoding)

    def render_simple(md_text: str, out_path: Path) -> int:
        doc = Document()
        # フォールバック時も可能なら游明朝に設定
        try:
            if qn is not None:
                styles = doc.styles
                for sname in [
                    "Normal",
                    "Heading 1", "Heading 2", "Heading 3",
                    "Heading 4", "Heading 5", "Heading 6",
                    "List Bullet", "List Number",
                ]:
                    try:
                        st = styles[sname]
                    except Exception:
                        continue
                    st.font.name = "游明朝"
                    try:
                        rPr = getattr(st._element, 'rPr', None)
                        if rPr is None:
                            rPr = st._element.get_or_add_rPr()
                        rFonts = getattr(rPr, 'rFonts', None)
                        if rFonts is None:
                            rFonts = rPr.get_or_add_rFonts()
                        rFonts.set(qn('w:eastAsia'), "游明朝")
                    except Exception:
                        pass
        except Exception:
            pass

        lines = md_text.splitlines()
        in_code = False
        code_lang = None
        para_count = 0

        bold_pat = re.compile(r"\*\*(.+?)\*\*")
        italic_pat = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
        inline_code_pat = re.compile(r"`([^`]+)`")

        def add_runs(paragraph, raw: str):
            # inline code → bold → italic の順で素朴に分割
            # 複雑なネストは未対応（最低限の整形）
            pos = 0
            for m in inline_code_pat.finditer(raw):
                head = raw[pos:m.start()]
                if head:
                    paragraph.add_run(head)
                code = m.group(1)
                r = paragraph.add_run(code)
                try:
                    r.font.name = 'Consolas'
                    r.font.size = Pt(10)
                except Exception:
                    pass
                pos = m.end()
            tail = raw[pos:]
            if not tail:
                return
            # bold
            parts = []
            last = 0
            for bm in bold_pat.finditer(tail):
                parts.append((False, tail[last:bm.start()]))
                parts.append(("bold", bm.group(1)))
                last = bm.end()
            parts.append((False, tail[last:]))
            # italic
            for kind, piece in parts:
                if kind == "bold":
                    r = paragraph.add_run(piece)
                    r.bold = True
                else:
                    s = piece
                    ilast = 0
                    for im in re.finditer(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", s):
                        t1 = s[ilast:im.start()]
                        if t1:
                            paragraph.add_run(t1)
                        r = paragraph.add_run(im.group(1))
                        r.italic = True
                        ilast = im.end()
                    t2 = s[ilast:]
                    if t2:
                        paragraph.add_run(t2)

        i = 0
        while i < len(lines):
            line = lines[i]
            # フェンスコード
            if line.strip().startswith("```"):
                if not in_code:
                    in_code = True
                    code_lang = line.strip().strip("`") or None
                else:
                    in_code = False
                    code_lang = None
                i += 1
                continue
            if in_code:
                p = doc.add_paragraph()
                r = p.add_run(line)
                try:
                    r.font.name = 'Consolas'
                    r.font.size = Pt(10)
                except Exception:
                    pass
                para_count += 1
                i += 1
                continue

            # 見出し
            m = re.match(r"^(#{1,6})\s+(.*)$", line)
            if m:
                level = len(m.group(1))
                text = m.group(2).strip()
                style = {
                    1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3',
                    4: 'Heading 4', 5: 'Heading 5', 6: 'Heading 6',
                }.get(level, 'Heading 3')
                p = doc.add_paragraph(style=style)
                add_runs(p, text)
                para_count += 1
                i += 1
                continue

            # 箇条書き（- or *）
            if re.match(r"^\s*[-*]\s+", line):
                text = re.sub(r"^\s*[-*]\s+", "", line)
                p = doc.add_paragraph(style='List Bullet')
                add_runs(p, text)
                para_count += 1
                i += 1
                continue

            # 番号付き
            if re.match(r"^\s*\d+\.\s+", line):
                text = re.sub(r"^\s*\d+\.\s+", "", line)
                p = doc.add_paragraph(style='List Number')
                add_runs(p, text)
                para_count += 1
                i += 1
                continue

            # 空行 → 段落区切り
            if not line.strip():
                i += 1
                continue

            # 通常段落
            p = doc.add_paragraph()
            add_runs(p, line)
            para_count += 1
            i += 1

        doc.save(str(out_path))
        return para_count

    count = render_simple(text, dst)
    return (str(dst), count)
